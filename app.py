"""
DocBlaze Backend - OpenAI Version with Feedback System (Phase 1 & 2)
Production-Grade RAG with Intelligent Feedback Loop
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import List, Optional, Dict, Any
import os
import uuid
from datetime import datetime, timezone
import json
from pathlib import Path
import hashlib
from docx import Document as DocxDocument

# MongoDB
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo import DESCENDING

# Document processing
import PyPDF2
from io import BytesIO
from PIL import Image
import pandas as pd
import easyocr
import cv2
import numpy as np
from pdf2image import convert_from_bytes

# LangChain with OpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate

# OpenAI direct import
from openai import OpenAI

# CrossEncoder for reranking
from sentence_transformers import CrossEncoder

# Initialize FastAPI
app = FastAPI(title="DocBlaze API - OpenAI with Feedback", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
VECTOR_STORE_DIR = Path("vector_stores")
VECTOR_STORE_DIR.mkdir(exist_ok=True)

MONGO_URL = os.getenv("MONGO_URL", "mongodb+srv://hritickra99_db_user:wXx0JRLwX1Kq8GP4@cluster0.r1fszik.mongodb.net/?appName=Cluster0")
OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY", "")
DB_NAME = "docblaze"

if not OPENAI_API_KEY:
    print("⚠️ WARNING: OPENAI_API_KEY not set!")
    print("Set it with: $env:OPENAI_API_KEY='your-key-here'")

# MongoDB client
mongo_client = AsyncIOMotorClient(MONGO_URL)
db = mongo_client[DB_NAME]

# Collections
users_collection = db.users
sessions_collection = db.sessions
documents_collection = db.documents
chat_history_collection = db.chat_history

# Initialize OpenAI
embeddings = OpenAIEmbeddings(
    model="text-embedding-3-small",
    openai_api_key=OPENAI_API_KEY
)

openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize CrossEncoder for reranking
# Initialize PaddleOCR (runs once at startup)
# Initialize EasyOCR (runs once at startup)
# Supports 80+ languages, loads model on first use
print("🔄 Initializing EasyOCR...")
easy_reader = easyocr.Reader(['en'], gpu=False)  # Set gpu=True if you have CUDA
print("✅ EasyOCR initialized successfully!")
cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

# Helper Functions

import numpy as np

def mongo_safe(obj):
    if isinstance(obj, dict):
        return {k: mongo_safe(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [mongo_safe(v) for v in obj]
    elif isinstance(obj, np.generic):
        return obj.item()
    else:
        return obj


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

async def verify_user_token(authorization: str = Header(None)) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization header")
    token = authorization.replace("Bearer ", "")
    user = await users_collection.find_one({"user_id": token})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid token")
    return token

# Document Processing
def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text = ""

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        if text.strip():
            return text.strip()

        # EasyOCR fallback for scanned PDFs
        print("📄 PDF has no extractable text, using EasyOCR...")
        images = convert_from_bytes(file_bytes, dpi=300)  # Higher DPI for better quality
        ocr_text = ""
        
        for i, img in enumerate(images):
            print(f"   Processing page {i+1}/{len(images)}...")
            
            # Preprocessing
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize if needed
            width, height = img.size
            if width < 1200:
                scale_factor = 1200 / width
                img = img.resize((int(width * scale_factor), int(height * scale_factor)), Image.LANCZOS)
            
            img_array = np.array(img)
            
            # EasyOCR with better parameters
            results = easy_reader.readtext(
                img_array,
                detail=1,
                paragraph=False,
                text_threshold=0.6,
                low_text=0.3
            )
            
            # Sort and extract text
            results_sorted = sorted(results, key=lambda x: (x[0][0][1], x[0][0][0]))
            
            for (bbox, text, confidence) in results_sorted:
                if confidence > 0.2:
                    ocr_text += text + " "
            
            ocr_text += "\n\n"

        return ocr_text.strip()

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_table(file_bytes: bytes, ext: str) -> str:
    if ext == "csv":
        df = pd.read_csv(BytesIO(file_bytes))
    else:
        df = pd.read_excel(BytesIO(file_bytes))

    return df.to_csv(index=False)

def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from image using EasyOCR with preprocessing"""
    img = Image.open(BytesIO(file_bytes))
    
    # Preprocessing for better OCR accuracy
    # Convert to RGB if needed
    if img.mode != 'RGB':
        img = img.convert('RGB')
    
    # Resize if too small (helps with ID cards)
    width, height = img.size
    if width < 800:
        scale_factor = 1200 / width
        new_width = int(width * scale_factor)
        new_height = int(height * scale_factor)
        img = img.resize((new_width, new_height), Image.LANCZOS)
    
    img_array = np.array(img)
    
    # EasyOCR returns list of (bbox, text, confidence)
    results = easy_reader.readtext(
        img_array,
        detail=1,  # Return confidence scores
        paragraph=False,  # Don't merge lines
        min_size=10,  # Minimum text box size
        text_threshold=0.6,  # Lower threshold for better recall
        low_text=0.3
    )
    
    # Extract text with confidence filtering and spatial ordering
    # Sort by vertical position (top to bottom, left to right)
    results_sorted = sorted(results, key=lambda x: (x[0][0][1], x[0][0][0]))
    
    text_lines = []
    prev_y = 0
    current_line = []
    
    for (bbox, detected_text, confidence) in results_sorted:
        if confidence > 0.2:  # Lower threshold for IDs
            # Get vertical position
            y_pos = bbox[0][1]
            
            # Check if same line (within 20 pixels)
            if prev_y == 0 or abs(y_pos - prev_y) < 20:
                current_line.append(detected_text)
            else:
                # New line detected
                if current_line:
                    text_lines.append(" ".join(current_line))
                current_line = [detected_text]
            
            prev_y = y_pos
    
    # Add last line
    if current_line:
        text_lines.append(" ".join(current_line))
    
    # Join lines with newlines
    extracted_text = "\n".join(text_lines)
    
    # Post-processing for Aadhaar numbers
    # Ensure Aadhaar numbers are on same line
    import re
    extracted_text = re.sub(r'(\d{4})\s*\n\s*(\d{4})\s*\n\s*(\d{4})', r'\1 \2 \3', extracted_text)
    
    print(f"📝 OCR extracted {len(text_lines)} lines with average confidence")
    
    return extracted_text.strip()

def extract_text(file_bytes: bytes, ext: str) -> str:
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext in ["xlsx", "xls", "csv"]:
        return extract_text_from_table(file_bytes, ext)
    elif ext in ["jpg", "jpeg", "png"]:
        return extract_text_from_image(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

# Vector Store Manager
import faiss

class VectorStoreManager:
    def __init__(self):
        self.stores: Dict[str, FAISS] = {}

    def get_store_path(self, user_id: str) -> Path:
        return VECTOR_STORE_DIR / f"user_{user_id}"

    async def get_or_create_store(self, user_id: str) -> FAISS:
        if user_id in self.stores:
            return self.stores[user_id]

        store_path = self.get_store_path(user_id)

        if store_path.exists():
            store = FAISS.load_local(
                str(store_path),
                embeddings,
                allow_dangerous_deserialization=True
            )
        else:
            # Proper empty FAISS index
            dim = len(embeddings.embed_query("test"))
            index = faiss.IndexFlatL2(dim)

            store = FAISS(
                embedding_function=embeddings,
                index=index,
                docstore={},
                index_to_docstore_id={}
            )

        self.stores[user_id] = store
        return store
    
    async def add_documents(self, user_id: str, texts: List[str], metadatas: List[Dict]):
        docs = [
            Document(page_content=text, metadata=meta)
            for text, meta in zip(texts, metadatas)
        ]
        
        store = await self.get_or_create_store(user_id)
        store.add_documents(docs)
        store.save_local(str(self.get_store_path(user_id)))

vector_manager = VectorStoreManager()

# 🔥 PHASE 2: LLM-as-Judge Evaluator
async def evaluate_response_quality(query: str, context: str, response: str) -> Dict[str, Any]:
    """Use LLM to evaluate response quality"""
    
    # Use more context for evaluation (increased from 1000 to 3000)
    context_sample = context[:3000] + "..." if len(context) > 3000 else context
    
    eval_prompt = f"""You are a quality evaluator. Analyze this RAG system response:

QUESTION: {query}

CONTEXT PROVIDED: {context_sample}

GENERATED RESPONSE: {response}

Evaluate on a scale of 0-1 for each metric:
1. FAITHFULNESS: Does the response accurately reflect the context? (0 = hallucination, 1 = fully faithful)
2. RELEVANCE: Does it answer the question? (0 = off-topic, 1 = directly answers)
3. COMPLETENESS: Is the answer complete? (0 = incomplete, 1 = comprehensive)

IMPORTANT: If you see specific data (like numbers, names, dates) in BOTH the context and response that match, score faithfulness high even if the exact phrasing differs.

Respond ONLY with valid JSON:
{{
    "faithfulness": 0.0-1.0,
    "relevance": 0.0-1.0,
    "completeness": 0.0-1.0,
    "overall_quality": 0.0-1.0,
    "concerns": ["list any issues"]
}}"""

    try:
        response_obj = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": eval_prompt}],
            temperature=0.1,
            max_tokens=300
        )
        
        result = json.loads(response_obj.choices[0].message.content)
        return result
    except Exception as e:
        print(f"LLM evaluation failed: {e}")
        return {
            "faithfulness": None,
            "relevance": None,
            "completeness": None,
            "overall_quality": None,
            "concerns": []
        }
        
# 🔥 PHASE 2: Track Implicit Feedback Signals
async def track_implicit_signals(session_id: str, chat_id: str, signal_type: str, value: Any):
    """Track implicit user behavior signals"""
    await db.implicit_signals.insert_one({
        "session_id": session_id,
        "chat_id": chat_id,
        "signal_type": signal_type,  # 'follow_up', 'explicit_feedback', etc.
        "value": value,
        "timestamp": datetime.now(timezone.utc)
    })

# Pydantic Models
class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    user_id: str
    username: str
    email: str
    created_at: str

class SessionResponse(BaseModel):
    session_id: str
    user_id: str
    session_name: str
    created_at: str
    last_active: str
    document_count: int

class QueryRequest(BaseModel):
    query: str
    regenerate: Optional[bool] = False
    previous_chat_id: Optional[str] = None

class QueryResponse(BaseModel):
    response: str
    sources: List[Dict[str, Any]]
    chat_id: str
    confidence_score: Optional[float] = None
    retrieval_quality: Optional[Dict[str, Any]] = None

class DocumentResponse(BaseModel):
    doc_id: str
    filename: str
    file_type: str
    extracted_text_preview: str
    uploaded_at: str

class FeedbackRequest(BaseModel):
    rating: str  # 'positive' or 'negative'
    comment: Optional[str] = None

class FeedbackResponse(BaseModel):
    chat_id: str
    rating: str
    timestamp: str

# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.post("/api/auth/register", response_model=UserResponse)
async def register_user(user: UserCreate):
    existing = await users_collection.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    hashed_password = hash_password(user.password)
    
    user_doc = {
        "user_id": user_id,
        "username": user.username,
        "email": user.email,
        "password": hashed_password,
        "created_at": datetime.now(timezone.utc)
    }
    
    await users_collection.insert_one(user_doc)
    
    return UserResponse(
        user_id=user_id,
        username=user.username,
        email=user.email,
        created_at=user_doc['created_at'].isoformat()
    )

@app.post("/api/auth/login")
async def login_user(credentials: UserLogin):
    user = await users_collection.find_one({"email": credentials.email})
    
    if not user or user['password'] != hash_password(credentials.password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    return {
        "token": user['user_id'],
        "user": UserResponse(
            user_id=user['user_id'],
            username=user['username'],
            email=user['email'],
            created_at=user['created_at'].isoformat()
        )
    }

@app.get("/api/users/me", response_model=UserResponse)
async def get_current_user(user_id: str = Depends(verify_user_token)):
    user = await users_collection.find_one({"user_id": user_id})
    
    return UserResponse(
        user_id=user['user_id'],
        username=user['username'],
        email=user['email'],
        created_at=user['created_at'].isoformat()
    )

@app.post("/api/sessions", response_model=SessionResponse)
async def create_session(
    session_name: str = "New Session",
    user_id: str = Depends(verify_user_token)
):
    session_id = str(uuid.uuid4())
    
    session_doc = {
        "session_id": session_id,
        "user_id": user_id,
        "session_name": session_name,
        "created_at": datetime.now(timezone.utc),
        "last_active": datetime.now(timezone.utc)
    }
    
    await sessions_collection.insert_one(session_doc)
    
    return SessionResponse(
        session_id=session_id,
        user_id=user_id,
        session_name=session_name,
        created_at=session_doc['created_at'].isoformat(),
        last_active=session_doc['last_active'].isoformat(),
        document_count=0
    )

@app.get("/api/sessions", response_model=List[SessionResponse])
async def get_user_sessions(user_id: str = Depends(verify_user_token)):
    sessions = await sessions_collection.find(
        {"user_id": user_id}
    ).sort("last_active", DESCENDING).to_list(length=100)
    
    result = []
    for session in sessions:
        doc_count = await documents_collection.count_documents(
            {"session_id": session['session_id']}
        )
        
        result.append(SessionResponse(
            session_id=session['session_id'],
            user_id=session['user_id'],
            session_name=session['session_name'],
            created_at=session['created_at'].isoformat(),
            last_active=session['last_active'].isoformat(),
            document_count=doc_count
        ))
    
    return result

@app.post("/api/sessions/{session_id}/documents", response_model=DocumentResponse)
async def upload_document(
    session_id: str,
    file: UploadFile = File(...),
    user_id: str = Depends(verify_user_token)
):
    # Validate session
    session = await sessions_collection.find_one(
        {"session_id": session_id, "user_id": user_id}
    )
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # Read file
    file_bytes = await file.read()
    file_extension = file.filename.split('.')[-1].lower()

    doc_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{doc_id}_{file.filename}"

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    # Extract text
    extracted_text = extract_text(file_bytes, file_extension)

    # Reject unreadable documents
    if not extracted_text or not extracted_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No readable text found in document. This may be a scanned or unsupported file."
        )

    # Store document metadata
    doc = {
        "doc_id": doc_id,
        "user_id": user_id,
        "session_id": session_id,
        "filename": file.filename,
        "file_path": str(file_path),
        "file_type": file_extension,
        "extracted_text": extracted_text,
        "uploaded_at": datetime.now(timezone.utc)
    }

    await documents_collection.insert_one(doc)

    # Chunk text for RAG
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_text(extracted_text)

    if not chunks:
        raise HTTPException(
            status_code=400,
            detail="Document could not be chunked into meaningful text."
        )

    metadatas = [
        {
            "doc_id": doc_id,
            "filename": file.filename,
            "file_type": file_extension,
            "session_id": session_id,
            "chunk_index": i
        }
        for i in range(len(chunks))
    ]

    # Add to FAISS
    await vector_manager.add_documents(user_id, chunks, metadatas)

    # Update session activity
    await sessions_collection.update_one(
        {"session_id": session_id},
        {"$set": {"last_active": datetime.now(timezone.utc)}}
    )

    return DocumentResponse(
        doc_id=doc_id,
        filename=file.filename,
        file_type=file_extension,
        extracted_text_preview=(
            extracted_text[:500] + "..."
            if len(extracted_text) > 500
            else extracted_text
        ),
        uploaded_at=doc["uploaded_at"].isoformat()
    )

@app.post("/api/sessions/{session_id}/query", response_model=QueryResponse)
async def query_documents(
    session_id: str,
    request: QueryRequest,
    user_id: str = Depends(verify_user_token)
):
    """
    🔥 ENHANCED QUERY WITH FEEDBACK INTELLIGENCE
    Phase 1: Feedback-based retrieval
    Phase 2: LLM-as-Judge + Implicit signals
    """
    session = await sessions_collection.find_one(
        {"session_id": session_id, "user_id": user_id}
    )
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Auto-rename session on first query
    if session.get("session_name") == "New Session":
        auto_name = request.query.strip()[:60]
        if len(request.query.strip()) > 60:
            auto_name += "..."
        await sessions_collection.update_one(
            {"session_id": session_id},
            {"$set": {"session_name": auto_name, "last_active": datetime.now(timezone.utc)}}
        )
    
    # 🔥 PHASE 2: Track follow-up query pattern (implicit signal)
    recent_chats = await chat_history_collection.find(
        {"session_id": session_id}
    ).sort("timestamp", -1).limit(1).to_list(length=1)
    
    if recent_chats:
        last_chat_time = recent_chats[0]['timestamp']
        
        # Make timezone-aware if it's naive
        if last_chat_time.tzinfo is None:
            last_chat_time = last_chat_time.replace(tzinfo=timezone.utc)
        
        time_diff = (datetime.now(timezone.utc) - last_chat_time).total_seconds()
        
        # If user asks another question within 2 minutes → potential quality issue
        if time_diff < 120:
            await track_implicit_signals(
                session_id, 
                str(recent_chats[0]['_id']), 
                'follow_up_query', 
                {'time_gap': time_diff, 'new_query': request.query}
            )
    
    # Get vector store
    vector_store = await vector_manager.get_or_create_store(user_id)
    
    # 🔥 PHASE 1: Feedback-based retrieval adjustment
    k_value = 20
    if request.regenerate:
        # Increase retrieval for regeneration
        k_value = 30
    
    # Retrieve documents
    retriever = vector_store.as_retriever(search_kwargs={"k": k_value})
    docs = retriever.invoke(request.query)
    
    if not docs:
        chat_id = str(uuid.uuid4())
        return QueryResponse(
            response="I couldn't find relevant information in your uploaded documents.",
            sources=[],
            chat_id=chat_id,
            confidence_score=0.0
        )
    
    # 🔥 PHASE 1: Get feedback history to boost/demote sources
    feedback_history = await chat_history_collection.find(
        {"session_id": session_id, "feedback.rating": {"$exists": True}}
    ).to_list(length=100)
    
    # Create source quality map based on feedback
    source_quality = {}
    for chat in feedback_history:
        if chat.get('sources'):
            for source in chat['sources']:
                doc_id = source.get('doc_id')
                if doc_id:
                    rating = chat.get('feedback', {}).get('rating')
                    if rating == 'positive':
                        source_quality[doc_id] = source_quality.get(doc_id, 0) + 0.1
                    elif rating == 'negative':
                        source_quality[doc_id] = source_quality.get(doc_id, 0) - 0.1
    
    # Rerank using CrossEncoder
    pairs = [(request.query, doc.page_content) for doc in docs]
    scores = cross_encoder.predict(pairs)
    
    # 🔥 PHASE 1: Apply feedback-based boost
    adjusted_scores = []
    for i, (doc, score) in enumerate(zip(docs, scores)):
        doc_id = doc.metadata.get('doc_id')
        boost = source_quality.get(doc_id, 0.0)
        adjusted_score = score + boost
        adjusted_scores.append(adjusted_score)
    
    # Sort docs by adjusted scores
    docs_with_scores = list(zip(docs, adjusted_scores))
    docs_with_scores.sort(key=lambda x: x[1], reverse=True)
    
    # Take top 5
    top_docs = [doc for doc, score in docs_with_scores[:5]]
    top_scores = [score for doc, score in docs_with_scores[:5]]
    
    # Calculate confidence score
    avg_confidence = sum(top_scores) / len(top_scores) if top_scores else 0.0
    # Normalize to 0-1 range (cross-encoder scores are typically -10 to 10)
    normalized_confidence = max(0.0, min(1.0, (avg_confidence + 10) / 20))
    
    # Smart context building - track which docs are actually used
    context_parts = []
    used_doc_ids = set()  # Track which documents actually contributed to context
    doc_relevance_scores = {}  # ✅ NEW: Track relevance score for each doc
    total_chars = 0
    MAX_CONTEXT_CHARS = 16000

    for doc, score in zip(top_docs, top_scores):  # ✅ CHANGED: Now iterates with scores
        chunk = doc.page_content
        doc_id = doc.metadata.get("doc_id")
        
        if total_chars + len(chunk) < MAX_CONTEXT_CHARS:
            context_parts.append(chunk)
            used_doc_ids.add(doc_id)
            
            # ✅ NEW: Track the best (highest) score for each document
            if doc_id not in doc_relevance_scores or score > doc_relevance_scores[doc_id]:
                doc_relevance_scores[doc_id] = score
            
            total_chars += len(chunk)
        else:
            break

    context = "\n\n---\n\n".join(context_parts)
    
    # ✅ NEW: Calculate relevance threshold - only show highly relevant sources
    # Use mean + 0.5*std deviation, or top 60% of scores
    if doc_relevance_scores:
        scores_list = list(doc_relevance_scores.values())
        avg_score = sum(scores_list) / len(scores_list)
        
        # Calculate standard deviation
        variance = sum((s - avg_score) ** 2 for s in scores_list) / len(scores_list)
        std_dev = variance ** 0.5
        
        # Threshold: mean - 0.3*std (keeps documents above average)
        relevance_threshold = avg_score - 0.3 * std_dev
        
        # Filter out low-relevance documents
        filtered_doc_ids = {
            doc_id for doc_id, score in doc_relevance_scores.items()
            if score >= relevance_threshold
        }
    else:
        filtered_doc_ids = used_doc_ids
    
    print(f"\n📋 CONTEXT STATS:")
    print(f"   - Total context length: {len(context)} chars")
    print(f"   - Chunks used: {len(context_parts)}/{len(top_docs)}")
    print(f"   - Documents contributing: {len(used_doc_ids)}")
    print(f"   - Relevance threshold: {relevance_threshold:.3f}")
    print(f"   - High-relevance docs: {len(filtered_doc_ids)}")
    
    # Get recent chat history
    history = await chat_history_collection.find(
        {"session_id": session_id}
    ).sort("timestamp", -1).limit(3).to_list(length=3)
    
    history.reverse()
    history_text = "\n".join([
        f"User: {h['query']}\nAssistant: {h['response'][:200]}..." 
        for h in history
    ])
    
    # Enhanced prompt
    prompt = f"""You are DocBlaze AI, an expert document analysis assistant.

CONTEXT from uploaded documents:
{context}

CONVERSATION HISTORY:
{history_text}

USER QUESTION:
{request.query}

INSTRUCTIONS:
1. Answer ONLY using information from the CONTEXT above
2. If exact numbers are asked, state clearly if calculation isn't possible
3. Cite specific document sections when relevant
4. If the answer isn't in the context, say: "This information is not available in the uploaded documents"
5. Be concise but complete
6. {"[REGENERATION MODE: Provide more detailed analysis]" if request.regenerate else ""}

ANSWER:"""

    # Call OpenAI
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2 if not request.regenerate else 0.3,
            max_tokens=800
        )
        
        answer = response.choices[0].message.content
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI API error: {str(e)}")
    
    # Build deduplicated sources - ONLY from documents actually used in context
    # Build deduplicated sources - ONLY from highly relevant documents
    sources_map = {}
    for doc, score in zip(top_docs, top_scores):  # ✅ CHANGED: Iterate with scores
        doc_id = doc.metadata.get("doc_id")
        
        # ⚠️ CRITICAL: Only include if this doc:
        # 1. Contributed to the context
        # 2. Has high relevance score (above threshold)
        if doc_id not in filtered_doc_ids:  # ✅ CHANGED: Use filtered_doc_ids instead of used_doc_ids
            continue
            
        filename = doc.metadata.get("filename", "Unknown")
        
        if doc_id and doc_id not in sources_map:
            sources_map[doc_id] = {
                "doc_id": doc_id,
                "filename": filename,
                "file_type": doc.metadata.get("file_type", "unknown"),
                "preview": doc.page_content[:200] + "...",
                "relevance_score": float(doc_relevance_scores.get(doc_id, 0))  # ✅ NEW: Add score
            }

    sources = list(sources_map.values())

    # Additional deduplication by filename (keep first occurrence)
    seen_filenames = set()
    unique_sources = []
    for source in sources:
        filename_lower = source["filename"].lower()
        if filename_lower not in seen_filenames:
            seen_filenames.add(filename_lower)
            unique_sources.append(source)

    # ✅ NEW: Sort by relevance score (highest first)
    unique_sources.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)

    # ✅ NEW: Remove internal relevance_score from API response (keep for debugging in logs)
    for source in unique_sources:
        source.pop('relevance_score', None)

    print(f"   - Unique sources shown: {len(unique_sources)}")
    if unique_sources:
        print("   - Source files with scores:")
        for source in unique_sources:
            score = doc_relevance_scores.get(source['doc_id'], 0)
            print(f"      • {source['filename']}: {score:.3f}") # 🔥 PHASE 2: LLM-as-Judge evaluation
    llm_evaluation = await evaluate_response_quality(request.query, context, answer)
    
    # Calculate retrieval quality metrics
    retrieval_quality = {
        "avg_cross_encoder_score": float(avg_confidence),
        "normalized_confidence": float(normalized_confidence),
        "num_sources_used": len(unique_sources),
        "context_diversity": len(set(s['doc_id'] for s in unique_sources)),
        "chunks_retrieved": len(top_docs),  # ✅ NEW: Total chunks retrieved
        "chunks_used_in_context": len(context_parts),  # ✅ NEW: Chunks actually used
        "context_length_chars": len(context),  # ✅ NEW: Context size
        "llm_evaluation": llm_evaluation
    }
    
    # Save to database
    chat_id = str(uuid.uuid4())
    chat_doc = {
        "chat_id": chat_id,
        "session_id": session_id,
        "user_id": user_id,
        "query": request.query,
        "response": answer,
        "sources": unique_sources,
        "confidence_score": normalized_confidence,
        "retrieval_quality": retrieval_quality,
        "timestamp": datetime.now(timezone.utc),
        "regenerated_from": request.previous_chat_id if request.regenerate else None
    }
    
    safe_chat_doc = mongo_safe(chat_doc)
    result = await chat_history_collection.insert_one(safe_chat_doc)

    print(f"✅ Saved chat with chat_id: {chat_id}, MongoDB _id: {result.inserted_id}")
    
    # Update session
    await sessions_collection.update_one(
        {"session_id": session_id},
        {"$set": {"last_active": datetime.now(timezone.utc)}}
    )
    
    return QueryResponse(
        response=answer,
        sources=unique_sources,
        chat_id=chat_id,
        confidence_score=normalized_confidence,
        retrieval_quality=retrieval_quality
    )

@app.get("/api/sessions/{session_id}/history")
async def get_session_history(
    session_id: str,
    user_id: str = Depends(verify_user_token)
):
    history = await chat_history_collection.find(
        {"session_id": session_id, "user_id": user_id}
    ).sort("timestamp", 1).to_list(length=500)
    
    return {
        "session_id": session_id,
        "history": [
            {
                "chat_id": chat['chat_id'],
                "query": chat['query'],
                "response": chat['response'],
                "sources": chat.get('sources', []),
                "confidence_score": chat.get('confidence_score'),
                "feedback": chat.get('feedback'),
                "timestamp": chat['timestamp'].isoformat()
            }
            for chat in history
        ]
    }

@app.get("/api/sessions/{session_id}/documents")
async def get_session_documents(
    session_id: str,
    user_id: str = Depends(verify_user_token)
):
    docs = await documents_collection.find(
        {"session_id": session_id, "user_id": user_id}
    ).sort("uploaded_at", DESCENDING).to_list(length=100)
    
    return {
        "session_id": session_id,
        "documents": [
            {
                "doc_id": doc['doc_id'],
                "filename": doc['filename'],
                "file_type": doc['file_type'],
                "uploaded_at": doc['uploaded_at'].isoformat()
            }
            for doc in docs
        ]
    }

# ============================================================================
# 🔥 NEW FEEDBACK ENDPOINTS
# ============================================================================

@app.post("/api/chats/{chat_id}/feedback", response_model=FeedbackResponse)
async def submit_feedback(
    chat_id: str,
    feedback: FeedbackRequest,
    user_id: str = Depends(verify_user_token)
):
    """Submit user feedback on a response"""
    
    # Validate rating
    if feedback.rating not in ['positive', 'negative']:
        raise HTTPException(status_code=400, detail="Rating must be 'positive' or 'negative'")
    
    print(f"🔍 Looking for chat_id: {chat_id}, user_id: {user_id}")
    
    # Find chat by chat_id field
    chat = await chat_history_collection.find_one({"chat_id": chat_id, "user_id": user_id})
    
    if not chat:
        print(f"❌ Chat not found: chat_id={chat_id}")
        # Debug: show recent chat_ids
        all_chats = await chat_history_collection.find({"user_id": user_id}).limit(5).to_list(length=5)
        print(f"📋 Recent chat_ids: {[c.get('chat_id') for c in all_chats]}")
        raise HTTPException(status_code=404, detail=f"Chat not found with chat_id: {chat_id}")
    
    print(f"✅ Found chat: {chat.get('query', 'No query')[:50]}...")
    
    # Update chat with feedback
    result = await chat_history_collection.update_one(
        {"chat_id": chat_id, "user_id": user_id},
        {
            "$set": {
                "feedback": {
                    "rating": feedback.rating,
                    "comment": feedback.comment,
                    "timestamp": datetime.now(timezone.utc)
                }
            }
        }
    )
    
    print(f"✅ Feedback saved: {feedback.rating}")
    
    # Track implicit signal
    await track_implicit_signals(
        chat['session_id'],
        chat_id,
        'explicit_feedback',
        {'rating': feedback.rating, 'comment': feedback.comment}
    )
    
    # ✅ ADD THIS RETURN STATEMENT
    return FeedbackResponse(
        chat_id=chat_id,
        rating=feedback.rating,
        timestamp=datetime.now(timezone.utc).isoformat()
    )
    
@app.get("/api/sessions/{session_id}/analytics")
async def get_session_analytics(
    session_id: str,
    user_id: str = Depends(verify_user_token)
):
    """Get analytics for a session"""
    
    # Verify session belongs to user
    session = await sessions_collection.find_one(
        {"session_id": session_id, "user_id": user_id}
    )
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all chats for session
    chats = await chat_history_collection.find(
        {"session_id": session_id}
    ).to_list(length=1000)
    
    total_queries = len(chats)
    
    # Count feedback
    positive_feedback = sum(1 for c in chats if c.get('feedback', {}).get('rating') == 'positive')
    negative_feedback = sum(1 for c in chats if c.get('feedback', {}).get('rating') == 'negative')
    
    # Calculate average confidence
    confidence_scores = [c.get('confidence_score') for c in chats if c.get('confidence_score') is not None]
    avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else None
    
    # Get LLM quality metrics
    llm_quality_scores = []
    for chat in chats:
        quality = chat.get('retrieval_quality', {}).get('llm_evaluation', {})
        if quality.get('overall_quality') is not None:
            llm_quality_scores.append(quality['overall_quality'])
    
    avg_llm_quality = sum(llm_quality_scores) / len(llm_quality_scores) if llm_quality_scores else None
    
    # Implicit signals analysis
    signals = await db.implicit_signals.find(
        {"session_id": session_id}
    ).to_list(length=1000)
    
    follow_up_count = sum(1 for s in signals if s.get('signal_type') == 'follow_up_query')
    
    return {
        "session_id": session_id,
        "total_queries": total_queries,
        "positive_feedback": positive_feedback,
        "negative_feedback": negative_feedback,
        "feedback_rate": (positive_feedback + negative_feedback) / total_queries if total_queries > 0 else 0,
        "avg_confidence": avg_confidence,
        "avg_llm_quality": avg_llm_quality,
        "follow_up_queries": follow_up_count,
        "quality_distribution": {
            "high": sum(1 for c in chats if c.get('confidence_score', 0) >= 0.7),
            "medium": sum(1 for c in chats if 0.4 <= c.get('confidence_score', 0) < 0.7),
            "low": sum(1 for c in chats if c.get('confidence_score', 0) < 0.4)
        }
    }

@app.get("/api/admin/feedback-dashboard")
async def get_feedback_dashboard(
    user_id: str = Depends(verify_user_token),
    limit: int = 50
):
    """Admin dashboard for reviewing feedback"""
    
    # Get recent negative feedback
    negative_chats = await chat_history_collection.find(
        {"user_id": user_id, "feedback.rating": "negative"}
    ).sort("timestamp", -1).limit(limit).to_list(length=limit)
    
    # Get low confidence responses
    low_confidence = await chat_history_collection.find(
        {"user_id": user_id, "confidence_score": {"$lt": 0.3}}
    ).sort("timestamp", -1).limit(limit).to_list(length=limit)
    
    return {
        "negative_feedback": [
            {
                "chat_id": chat['chat_id'],
                "session_id": chat['session_id'],
                "query": chat['query'],
                "response": chat['response'][:200] + "...",
                "confidence_score": chat.get('confidence_score'),
                "feedback_comment": chat.get('feedback', {}).get('comment'),
                "timestamp": chat['timestamp'].isoformat()
            }
            for chat in negative_chats
        ],
        "low_confidence_responses": [
            {
                "chat_id": chat['chat_id'],
                "session_id": chat['session_id'],
                "query": chat['query'],
                "response": chat['response'][:200] + "...",
                "confidence_score": chat.get('confidence_score'),
                "timestamp": chat['timestamp'].isoformat()
            }
            for chat in low_confidence
        ]
    }
    
@app.get("/api/documents/{doc_id}/preview")
async def get_document_preview(
    doc_id: str,
    user_id: str = Depends(verify_user_token)
):
    """Debug endpoint: Get full extracted text from a document"""
    doc = await documents_collection.find_one({
        "doc_id": doc_id,
        "user_id": user_id
    })
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "doc_id": doc_id,
        "filename": doc['filename'],
        "file_type": doc['file_type'],
        "extracted_text": doc['extracted_text'],
        "text_length": len(doc['extracted_text']),
        "uploaded_at": doc['uploaded_at'].isoformat()
    }

    
from pathlib import Path
from fastapi.staticfiles import StaticFiles
BASE_DIR = Path(__file__).resolve().parent
app.mount(
    "/",
    StaticFiles(directory=BASE_DIR / "static", html=True),
    name="static"
)




@app.get("/health")
async def health_check():
    return {
        "status": "healthy", 
        "service": "DocBlaze API with Feedback Intelligence",
        "version": "2.0.0",
        "llm": "gpt-4o-mini",
        "features": [
            "Phase 1: Explicit Feedback & Reranking",
            "Phase 2: LLM-as-Judge & Implicit Signals"
        ]
    }

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting DocBlaze API with Feedback System...")
    print("📊 Features: Feedback Intelligence, LLM-as-Judge, Smart Reranking")
    uvicorn.run(app, host="0.0.0.0", port=5000)